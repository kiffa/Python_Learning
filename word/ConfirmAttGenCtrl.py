from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from trs_comm.models import TradeOrdr
from trs_comm.models import Cust
from trs_comm.ctrls.CntrCtrl import CntrCtrlObj
from trs_comm.ctrls.XferOrderCtrl import XferOrderCtrlObj
from datetime import datetime
from django.conf import settings
import os
import time
import mysite.tools.dbfuncs as dbfuncs

class ConfirmAttGenCtrl(object):

    def __init__(self):
        pass

    def page_content_handle(self, document, replace_content):
        #处理页眉
        for setcion in document.sections:
            paras = setcion.header.paragraphs
            for p in paras:
                for i in range(len(p.runs)):
                    for key, value in replace_content.items():
                        runs_text = p.runs[i].text
                        if runs_text == key or key in runs_text:
                            p.runs[i].text = runs_text.replace(key, value)

        # 处理段落
        for para in document.paragraphs:
            # 找到文本替换位置
            for i in range(len(para.runs)):
                for key, value in replace_content.items():
                    runs_text = para.runs[i].text
                    if runs_text == key or key in runs_text:
                        para.runs[i].text = runs_text.replace(key, value)
        return document

    #利率期货附件平仓
    def un_rate_att_table_handle(self,document,trade_info):
        trade_ordr_lst = trade_info.get('trade_ordr_lst')
        trade_summary = trade_info.get('trade_summary')
        trade_len = len(trade_ordr_lst)
        # 处理表格
        # 甲乙方信息
        table_0 = document.tables[0]
        for row in table_0.rows:  # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in trade_summary.items():
                    if key in cell.text:
                        for cell_para in cell.paragraphs:
                            for run in cell_para.runs:
                                if run.text == key:
                                    run.text = run.text.replace(run.text, value)

        # 标的平仓明细
        table_1 = document.tables[1]
        for row_idx, row in enumerate(table_1.rows):
            if row_idx > 0:  # 第一行位表头
                if row_idx - 1 < trade_len:
                    row.cells[0].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('CFM_DOC_ID')
                    row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('TRADE_TYPE')
                    row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('TRADE_DEAL_DATE')
                    row.cells[3].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('NAME')
                    row.cells[4].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('CODE')
                    row.cells[5].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('DACTB')
                    row.cells[6].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('DACT')
                    row.cells[7].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('DACTA')
                    row.cells[8].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('AMT')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)

        # 合约平仓信息
        table_2 = document.tables[2]
        for row_idx, row in enumerate(table_2.rows):
            if row_idx > 0:  # 第一行位表头
                if row_idx - 1 < trade_len:
                    row.cells[0].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('CFM_DOC_ID')
                    row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('TRADE_DEAL_DATE')
                    row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('DEACOST')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)
        return document

    # 利率期货开仓附件
    def rate_att_table_handle(self, document, trade_info):
        trade_ordr_lst = trade_info.get('trade_ordr_lst')
        trade_summary = trade_info.get('trade_summary')
        trade_len = len(trade_ordr_lst)
        # 处理表格
        # 甲乙方信息
        table_0 = document.tables[0]
        for row in table_0.rows:  # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in trade_summary.items():
                    if key in cell.text:
                        for cell_para in cell.paragraphs:
                            for run in cell_para.runs:
                                if run.text == key:
                                    run.text = run.text.replace(run.text, value)
        # 基本要素
        table_2 = document.tables[2]
        t2_rows = table_2.rows
        t2_rows[0].cells[1].paragraphs[0].runs[0].text = trade_summary['TRADE_DEAL_DATE']  # 交易达成日
        t2_rows[0].cells[3].paragraphs[0].runs[0].text = trade_summary['TRADE_ORDER_DATE']  # 起始日
        t2_rows[1].cells[1].paragraphs[0].runs[0].text = trade_summary['TRADE_FINISH_DATE']  # 到期日
        t2_rows[1].cells[3].paragraphs[0].runs[0].text = trade_summary['NOTIONAL_AMT_TOTAL']  # 合约名义本金额(交易货币)
        t2_rows[2].cells[1].paragraphs[0].runs[0].text = trade_summary['PAYOFFCCY']  # 交易货币
        t2_rows[3].cells[1].paragraphs[0].runs[0].text = str(trade_summary['FEE_RATE'])  # 基本费率

        # 投资组合表格
        table_3 = document.tables[3]
        for row_idx, row in enumerate(table_3.rows):
            if row_idx > 0:  # 第一行位表头
                if row_idx - 1 < trade_len:
                    row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('NAME')
                    row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('CODE')
                    row.cells[3].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('AVGCOST')
                    row.cells[4].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('DEACOUNT')
                    row.cells[5].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('MULTIPLIER')
                    row.cells[6].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx - 1].get('TRADE_TYPE')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)
        return document
    def unconf_att_table_handle(self, document, trade_info):
        trade_ordr_lst = trade_info.get('trade_ordr_lst')
        trade_summary = trade_info.get('trade_summary')
        trade_len = len(trade_ordr_lst)
        # 处理表格
        # 甲乙方信息
        table_0 = document.tables[0]
        for row in table_0.rows:    # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in trade_summary.items():
                    if key in cell.text:
                        for cell_para in cell.paragraphs:
                            for run in cell_para.runs:
                                if run.text == key:
                                    run.text = run.text.replace(run.text, value)

        # 标的平仓明细
        table_1 = document.tables[1]
        for row_idx, row in enumerate(table_1.rows):
            if row_idx > 0: #第一行位表头
                if row_idx -1 < trade_len:
                    row.cells[0].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('CFM_DOC_ID')
                    row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('TRADE_TYPE')
                    row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('TRADE_DEAL_DATE')
                    row.cells[3].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('NAME')
                    row.cells[4].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('CODE')
                    row.cells[5].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('DACTB')
                    row.cells[6].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('DACT')
                    row.cells[7].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('DACTA')
                    row.cells[8].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('AMT')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)

        # 合约平仓信息
        table_2 = document.tables[2]
        for row_idx, row in enumerate(table_2.rows):
            if row_idx > 0: #第一行位表头
                if row_idx -1 < trade_len:
                    row.cells[0].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('CFM_DOC_ID')
                    row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('TRADE_DEAL_DATE')
                    row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('DEACOST')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)
        return document

    def conf_att_table_handle(self, document, trade_info):
        trade_ordr_lst = trade_info.get('trade_ordr_lst')
        trade_summary = trade_info.get('trade_summary')
        trade_len = len(trade_ordr_lst)
        # 处理表格
        # 甲乙方信息
        table_0 = document.tables[0]
        for row in table_0.rows:    # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in trade_summary.items():
                    if key in cell.text:
                        for cell_para in cell.paragraphs:
                            for run in cell_para.runs:
                                if run.text == key:
                                    run.text = run.text.replace(run.text, value)
        # 基本要素
        table_2 = document.tables[2]
        t2_rows = table_2.rows
        t2_rows[0].cells[1].paragraphs[0].runs[0].text = trade_summary['TRADE_DEAL_DATE'] # 交易达成日
        t2_rows[0].cells[3].paragraphs[0].runs[0].text = trade_summary['TRADE_ORDER_DATE'] # 起始日
        t2_rows[1].cells[1].paragraphs[0].runs[0].text = trade_summary['TRADE_FINISH_DATE'] # 到期日
        t2_rows[1].cells[3].paragraphs[0].runs[0].text = trade_summary['NOTIONAL_AMT_TOTAL'] # 合约名义本金额(交易货币)

        # 投资组合表格
        table_3 = document.tables[3]
        for row_idx, row in enumerate(table_3.rows):
            if row_idx > 0: #第一行位表头
                if row_idx -1 < trade_len:
                    row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('NAME')
                    row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('CODE')
                    row.cells[3].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('AVGCOST')
                    row.cells[4].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('DEACOUNT')
                    row.cells[5].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('MULTIPLIER')
                    row.cells[6].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx -1].get('TRADE_TYPE')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)

        # 履约保障比率
        table_4 = document.tables[4] # 履约保障比率
        for row_idx, row in enumerate(table_4.rows):
            if row_idx > 0: #第一行位表头
                if row_idx -1 < trade_len * 2:
                    if row_idx % 2 == 0:
                        row.cells[0].paragraphs[0].runs[0].text = trade_ordr_lst[int((row_idx - 2)/2)].get('NAME')
                else:  # 删除多余的列
                    row._element.getparent().remove(row._element)

        # 最后平仓日
        table_5 = document.tables[5] # 最后平仓日
        for row_idx, row in enumerate(table_5.rows):
            if row_idx < trade_len:    #
                row.cells[0].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx].get('NAME')
                row.cells[1].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx].get('CODE')
                row.cells[2].paragraphs[0].runs[0].text = trade_ordr_lst[row_idx].get('LAST_TRADE_DATE')
            else:  # 删除多余的列
                row._element.getparent().remove(row._element)
        return document

    # 商品单笔单签1
    def conf_1_att_table_handle(self, document, trade_info):
        trade_ordr_lst = trade_info.get('trade_ordr_lst')
        trade_ordr_info = trade_ordr_lst[0] if len(trade_ordr_lst) > 0 else {}
        trade_summary = trade_info.get('trade_summary')
        trade_summary.update(trade_ordr_info)
        # 处理表格
        # 业务要素及确认
        table_1 = document.tables[1]
        for row in table_1.rows:    # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in trade_summary.items():
                    if key == cell.text:
                        cell.text = str(value)
                    elif key in cell.text:
                        if 'FEE_RATE' == key:
                            pass
                        else:
                            cell.text = cell.text.replace(key, str(value))
        return document

    ## 开仓
    def gen_confirm_att(self, ORDER_CFM_DOC_ID):
        # 一、读取模板文件
        # 通过交易确认书编号查询对应合约类型
        cntr_type_sql = '''SELECT cntr.PRODUCT_TYPE_2,cntr.TRANS_TYPE,cntr.CURRENCY,cntr.BASIC_FEE_RATE,cntr.FILL_ROLE FROM `V_TRS_COMM_TRADE_ORDR` ordr, `t_trs_comm_cntr` cntr 
            WHERE ordr.`CONTRACT_ID` = cntr.`CONTRACT_ID` AND ordr.`ORDER_CFM_DOC_ID` = '{}'
        '''.format(ORDER_CFM_DOC_ID)
        cntr_type_data = dbfuncs.from_sql_manually(cntr_type_sql, DBConf="Project")
        cntr_type = cntr_type_data['PRODUCT_TYPE_2'].iloc[0]
        trans_type = cntr_type_data['TRANS_TYPE'].iloc[0]
        fill_role =  cntr_type_data['FILL_ROLE'].iloc[0]
        currency = cntr_type_data['CURRENCY'].iloc[0]
        fee_rate = cntr_type_data['BASIC_FEE_RATE'].iloc[0]
        eg_file_name = '跨境收益互换交易确认书-' + cntr_type + '(模板).docx'
        if trans_type and trans_type == '1':
            eg_file_name = '场外衍生品交易确认书-利率期货互换(模板).docx'
            if fill_role and fill_role == '1':
                eg_file_name = '场外衍生品交易确认书-利率期货互换(乙方模板).docx'
        elif trans_type and trans_type == '2':
            eg_file_name = '跨境收益互换交易确认书-商品互换-单笔单签1(乙方模板).docx'
        file_path = os.path.join(settings.MEDIA_ROOT, 'eg', eg_file_name)
        document = Document(file_path)

        # 二、初始化数据
        sign_date = "" # 指令日期
        client = "" # 客户名称
        main_contract_id = "" # 主协议编号
        defn_doc_id = "" # 定义文件编号

        # 三、获取交易数据
        sql = '''SELECT t_order.TRADE_DEAL_DATE   TRADE_DEAL_DATE,
            t_order.TRADE_ORDER_DATE  TRADE_ORDER_DATE,
            t_order.TRADE_FINISH_DATE TRADE_FINISH_DATE,
            t_order.TRADE_TYPE        TRADE_TYPE,
            t_order.CONTRACT_ID       CONTRACT_ID,
            t_order.I_CODE            I_CODE,
            IFNULL(t_order.DEAAMOUNT, 0) DEAAMOUNT,
            IFNULL(t_order.DEAFEE, 0)    DEAFEE,
            IFNULL(t_order.AVGPRICE, 0)     AVGCOST,
            IFNULL(t_order.DEACOUNT, 0) OPEN_DEACOUNT,
            IFNULL(t_order.CAP_RATE, 0) CAP_RATE,
            IFNULL(t_order.CAP_LOC, 0) CAP_LOC,
            IFNULL(t_order.BASIC_FEE_RATE, 0) BASIC_FEE_RATE,
            sec.I_NAME                I_NAME,
            sec.MULTIPLIER            MULTIPLIER,
            sec.LAST_TRADE_DATE       LAST_TRADE_DATE,
            sec.I_CODE_BLOOMBERG      I_CODE_BLOOMBERG,
            IFNULL(sec.MARKET, sec.M_TYPE)      M_TYPE,
            IFNULL(marg.IM_RATE, sec.IM_RATE) IM_RATE,
            IFNULL(marg.VM_RATE, sec.VM_RATE) VM_RATE,
            cntr.DEFN_DOC_ID          DEFN_DOC_ID,
            cust.MAIN_CNTR_NAME	      MAIN_CNTR_NAME
        FROM V_TRS_COMM_TRADE_ORDR t_order
        LEFT JOIN T_TRS_COMM_SEC sec ON t_order.i_code = sec.i_code
        LEFT JOIN T_TRS_COMM_CNTR cntr ON t_order.CONTRACT_ID = cntr.CONTRACT_ID
        LEFT JOIN T_TRS_COMM_CUST cust ON cust.PARTY_ID = cntr.PARTY_ID
        LEFT JOIN T_TRS_COMM_MARG marg ON marg.CONTRACT_ID = cntr.CONTRACT_ID AND marg.i_code = sec.i_code
        where t_order.ORDER_CFM_DOC_ID = '{}'
        '''.format(ORDER_CFM_DOC_ID)
        TradeOrdr = dbfuncs.from_sql_manually(sql, DBConf="Project")
        MAIN_CNTR_NAME = TradeOrdr['MAIN_CNTR_NAME'][0]
        if not MAIN_CNTR_NAME:
            MAIN_CNTR_NAME = '中国证券期货市场场外衍生品交易主协议（2014版）'
        if len(TradeOrdr) > 0:
            TradeOrdr = TradeOrdr.fillna('')
            CONTRACT_ID = TradeOrdr.CONTRACT_ID.iloc[0]
            cntr_data_dic = CntrCtrlObj.list(CONTRACT_ID=CONTRACT_ID, rel=True)

            # 生成参数
            client = cntr_data_dic[0]['客户名称']
            main_contract_id = cntr_data_dic[0]['主协议编号']
            defn_doc_id = TradeOrdr.DEFN_DOC_ID.iloc[0]
            trade_info = self.trade_ordr_handle(TradeOrdr)
            trade_info['trade_summary']['CLIENT'] = client
            trade_info['trade_summary']['PAYOFFCCY'] = currency
            trade_info['trade_summary']['FEE_RATE'] = fee_rate
        # 四、生成待替换数据
        replace_para_content = {
            "CNTRID": ORDER_CFM_DOC_ID,
            "CLIENT":client,
            "MAIN_CONTRACT_ID":main_contract_id,
            "DEFN_DOC_ID":defn_doc_id,
            'MAIN_CNTR_NAME':MAIN_CNTR_NAME
        }

        # 五、处理文档
        # 1)、处理正文，页眉
        document = self.page_content_handle(document, replace_para_content)
        # 2)、处理数据表
        if trans_type and trans_type == '1':
            document = self.rate_att_table_handle(document, trade_info)
        elif trans_type and trans_type == '2':
            # 商品单笔单签1
            document = self.conf_1_att_table_handle(document, trade_info)
        else:
            document = self.conf_att_table_handle(document, trade_info)

        # 六、生成目标文件
        file_name = '场外衍生品交易确认书【' + ORDER_CFM_DOC_ID + '】.docx'
        if trans_type and trans_type == '1':
            file_name = '场外衍生品交易确认书【' + ORDER_CFM_DOC_ID + '】.docx'
        dest_temp_file = os.path.join(settings.MEDIA_ROOT, 'report', file_name)
        document.save(dest_temp_file)

        # 七、返回当前文件路径
        return {'dest_temp_file':dest_temp_file, 'file_name':file_name}

    ## 预付金变动
    def gen_yfjbd_doc(self,contract_id,date,jz_yfj):
        cntr_sql = "SELECT * FROM `t_trs_comm_cntr` cntr WHERE cntr.`CONTRACT_ID` = '{}'".format(contract_id)
        cntr_type_data = dbfuncs.from_sql_manually(cntr_sql, DBConf="Project")
        cntr_type = cntr_type_data['PRODUCT_TYPE_2'].iloc[0]
        trans_type = cntr_type_data['TRANS_TYPE'].iloc[0]
        fill_role = cntr_type_data['FILL_ROLE'].iloc[0]
        currency = '' if not cntr_type_data['CURRENCY'].iloc[0] else cntr_type_data['CURRENCY'].iloc[0]
        fee_rate = cntr_type_data['BASIC_FEE_RATE'].iloc[0]
        cust = Cust.objects.filter(PARTY_ID=cntr_type_data['PARTY_ID'].iloc[0])

        cntr_data_dic = CntrCtrlObj.list(CONTRACT_ID=contract_id, rel=True)
        MAIN_CNTR_NAME = '中国证券期货市场场外衍生品交易主协议'
        if cust.first().MAIN_CNTR_NAME:
            MAIN_CNTR_NAME = cust.first().MAIN_CNTR_NAME
        # 生成参数
        client = cntr_data_dic[0]['客户名称']
        main_contract_id = '' if not cntr_data_dic[0]['主协议编号'] else cntr_data_dic[0]['主协议编号']
        defn_doc_id = cntr_data_dic[0]['定义文件编号']
        date_time = datetime.strptime(date,"%Y-%m-%d")
        CLOSE_ORDER_CFM_DOC_ID = contract_id+'-CF'+ datetime.strftime(date_time,'%Y%m%d')
        sign_date = "{}年{}月{}日".format(str(date_time.year),str(date_time.month),str(date_time.day))  # 指令日期
        if trans_type and trans_type == '1':
            eg_file_name = "场外衍生品互换交易平仓及结算确认书（预付金变动）-利率期货(模板).docx"
            if fill_role and fill_role == '1':
                eg_file_name = "场外衍生品互换交易平仓及结算确认书（预付金变动）-利率期货(乙方模板).docx"
        else:
            eg_file_name = "场外衍生品互换交易平仓及结算确认书（预付金变动）-商品期货(模板).docx"
        # 四、生成待替换数据
        replace_para_content = {
            "CNTRID": CLOSE_ORDER_CFM_DOC_ID,
            "CLIENT": client,
            "MAIN_CONTRACT_ID": main_contract_id,
            "DEFN_DOC_ID": defn_doc_id,
            'MAIN_CNTR_NAME': MAIN_CNTR_NAME,
            'PAYOFFCCY': currency
        }
        xfer_data = dbfuncs.from_sql(table_name='T_TRS_COMM_XFER_ORDER',where="CONTRACT_ID='{}' and XFER_ORDER_DATE = '{}'".format(contract_id,date),DBConf='Project')
        XFER_BAL = 0
        CAP_LOC = 0
        PRE_LOC = 0
        # 获取出入金余额 保留两位小数 XFER_ORDER_TYPE
        if not xfer_data.empty:
            XFER_BAL = round(int(xfer_data['XFER_ORDER_TYPE'][0])*xfer_data['AMOUNT'][0],2)
        date_data = dbfuncs.get_last_trade_date(date_today=date)
        acct_data = dbfuncs.from_sql(table_name='T_TRS_COMM_ACCT_HIS',where="CONTRACT_ID='{}' and TRADE_DATE in ('{}','{}')".format(contract_id, date,str(date_data)),DBConf='Project')
        if not acct_data.empty:
            CAP_LOC = acct_data[acct_data['TRADE_DATE']==date]['CAP_LOC'].iloc[0]
            PRE_LOC = acct_data[acct_data['TRADE_DATE']==str(date_data)]['CAP_LOC'].iloc[0]
        table_content = {
            "CLIENT": client,
            'SETT_BAL': str(round(float(jz_yfj),2)),
            'SIGN_DATE': sign_date,
            'CHANGE_DATE': date,
            'XFER_BAL' : str(XFER_BAL),
            'CAP_LOC':str(round(float(CAP_LOC),2)),
            'PRE_LOC':str(round(float(PRE_LOC),2))
        }
        file_path = os.path.join(settings.MEDIA_ROOT, 'eg', eg_file_name)
        document = Document(file_path)
        document = self.page_content_handle(document, replace_para_content)
        # 处理表格
        # 甲乙方信息
        table_0 = document.tables[0]
        for row in table_0.rows:  # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in table_content.items():
                    if key in cell.text:
                        for cell_para in cell.paragraphs:
                            for run in cell_para.runs:
                                if run.text == key:
                                    run.text = run.text.replace(run.text, value)
        table_1 = document.tables[1]
        for row in table_1.rows:  # 遍历表格中的所有行
            for cell in row.cells:  # 遍历行中的所有单元格
                for key, value in table_content.items():
                    if key in cell.text:
                        for cell_para in cell.paragraphs:
                            for run in cell_para.runs:
                                if run.text == key:
                                    run.text = run.text.replace(run.text, value)

        file_name = '场外衍生品互换交易平仓及结算确认书（预付金变动）-【' + CLOSE_ORDER_CFM_DOC_ID + '】.docx'
        dest_temp_file = os.path.join(settings.MEDIA_ROOT, 'report', file_name)
        document.save(dest_temp_file)
        return {'dest_temp_file': dest_temp_file, 'file_name': file_name}
    ## 平仓
    def gen_un_confirm_att(self, CLOSE_ORDER_CFM_DOC_ID):
        # 一、读取模板文件
        # 通过交易确认书编号查询对应合约类型
        cntr_type_sql = '''SELECT cntr.PRODUCT_TYPE_2,cntr.TRANS_TYPE,cntr.CURRENCY,cntr.BASIC_FEE_RATE,cntr.FILL_ROLE FROM `T_TRS_COMM_TRADE_REC` trade, `t_trs_comm_cntr` cntr 
            WHERE trade.`CONTRACT_ID` = cntr.`CONTRACT_ID` AND trade.`CLOSE_ORDER_CFM_DOC_ID` = '{}'
        '''.format(CLOSE_ORDER_CFM_DOC_ID)
        cntr_type_data = dbfuncs.from_sql_manually(cntr_type_sql, DBConf="Project")
        cntr_type = cntr_type_data['PRODUCT_TYPE_2'].iloc[0]
        trans_type = cntr_type_data['TRANS_TYPE'].iloc[0]
        currency = '' if not cntr_type_data['CURRENCY'].iloc[0] else cntr_type_data['CURRENCY'].iloc[0]
        fill_role = cntr_type_data['FILL_ROLE'].iloc[0]
        fee_rate = cntr_type_data['BASIC_FEE_RATE'].iloc[0]
        eg_file_name = '场外衍生品互换交易平仓及结算确认书-' + cntr_type + '(模板).docx'
        if trans_type and trans_type == '1':
            eg_file_name = '场外衍生品交易平仓及结算确认书-利率期货互换(模板).docx'
            if fill_role and fill_role=='1':
                eg_file_name = '场外衍生品交易平仓及结算确认书-利率期货互换(乙方模板).docx'
        elif trans_type and trans_type == '2':
            eg_file_name = '场外衍生品互换交易平仓及结算确认书-商品互换-单笔单签1(乙方模板).docx'

        file_path = os.path.join(settings.MEDIA_ROOT, 'eg', eg_file_name)
        document = Document(file_path)

        # 二、初始化数据
        sign_date = "" # 指令日期
        client = "" # 客户名称
        main_contract_id = "" # 主协议编号
        defn_doc_id = "" # 定义文件编号        

        # 三、获取交易数据
        sql = '''SELECT T_TRADE.CLOSE_ORDER_CFM_DOC_ID,
            T_TRADE.OPEN_ORDER_CFM_DOC_ID,
            T_TRADE.CLOSE_DATE,
            T_TRADE.I_CODE,
            SEC.I_NAME,
            SEC.I_CODE_BLOOMBERG,
            T_TRADE.TRADE_TYPE,
            T_TRADE.DEACOUNT,
            IFNULL(T_POS.DEACOUNT, 0) DEACOUNT_POS,
            T_TRADE.CLOSE_DEAAMOUNT,
            T_TRADE.CLOSE_DEAFEE,
            IFNULL(T_POS.DEACOST, 0) DEACOST_POS,
            T_TRADE.CONTRACT_ID
        FROM T_TRS_COMM_TRADE_REC T_TRADE
        LEFT JOIN T_TRS_COMM_POSI_HIS T_POS
            ON T_POS.OPEN_ORDER_CFM_DOC_ID = T_TRADE.OPEN_ORDER_CFM_DOC_ID
        AND T_TRADE.CLOSE_DATE = T_POS.TRADE_DATE
        AND T_POS.I_CODE = T_TRADE.I_CODE
        LEFT JOIN T_TRS_COMM_SEC SEC
            ON T_TRADE.I_CODE = SEC.I_CODE
        WHERE T_TRADE.CLOSE_ORDER_CFM_DOC_ID = '{}'
        '''.format(CLOSE_ORDER_CFM_DOC_ID)
        trade_data = dbfuncs.from_sql_manually(sql, DBConf="Project")
        MAIN_CNTR_NAME = '中国证券期货市场场外衍生品交易主协议（2014版）'
        if len(trade_data) > 0:
            CONTRACT_ID = trade_data.CONTRACT_ID.loc[0]
            cntr_data_dic = CntrCtrlObj.list(CONTRACT_ID=CONTRACT_ID, rel=True)
            cust = Cust.objects.filter(PARTY_ID=cntr_data_dic[0]['客户编号'])
            if cust.first().MAIN_CNTR_NAME:
                MAIN_CNTR_NAME = cust.first().MAIN_CNTR_NAME
            # 生成参数
            client = cntr_data_dic[0]['客户名称']
            main_contract_id ='' if not cntr_data_dic[0]['主协议编号'] else cntr_data_dic[0]['主协议编号']
            defn_doc_id = cntr_data_dic[0]['定义文件编号']
            trade_info = self.trade_unordr_handle(trade_data)
            trade_info['trade_summary']['CLIENT'] = client

        # 四、生成待替换数据
        replace_para_content = {
            "CNTRID": CLOSE_ORDER_CFM_DOC_ID,
            "CLIENT":client,
            "MAIN_CONTRACT_ID":main_contract_id,
            "DEFN_DOC_ID":defn_doc_id,
            'MAIN_CNTR_NAME':MAIN_CNTR_NAME,
            'PAYOFFCCY': currency
        }

        # 五、处理文档
        # 1)、处理正文，页眉
        document = self.page_content_handle(document, replace_para_content)
        # 2)、处理数据表
        if trans_type and trans_type == '1':
            document = self.un_rate_att_table_handle(document, trade_info)
        else:
            document = self.unconf_att_table_handle(document, trade_info)

        # 六、生成目标文件
        file_name = '场外衍生品交易平仓及结算确认书-【' + CLOSE_ORDER_CFM_DOC_ID + '】.docx'
        dest_temp_file = os.path.join(settings.MEDIA_ROOT, 'report', file_name)
        document.save(dest_temp_file)

        # 七、返回当前文件路径
        return {'dest_temp_file':dest_temp_file, 'file_name':file_name}

    def trade_ordr_handle(self, TradeOrdr):
        trade_ordr_lst = []
        trade_summary = {}

        TRADE_DEAL_DATE = "" # 交易达成日
        TRADE_ORDER_DATE = "" # 起始日
        TRADE_FINISH_DATE = "" # 到期日
        NOTIONAL_AMT_TOTAL = 0 # 合约名义本金
        SIGN_DATE = "" # 日期
        for trade_idx, trade_order in TradeOrdr.iterrows():
            if trade_idx == 0:
                TRADE_DEAL_DATE = trade_order['TRADE_DEAL_DATE']
                TRADE_ORDER_DATE = trade_order['TRADE_ORDER_DATE']
                trade_date_lst = TRADE_ORDER_DATE.split('-')
                if trade_date_lst:
                    SIGN_DATE = trade_date_lst[0] + ' 年 ' + trade_date_lst[1] + ' 月 ' + trade_date_lst[2] + ' 日'
            I_CODE = "" # 标的代码
            I_NAME = "" # 标的名称
            LAST_TRADE_DATE = "" # 最后平仓日
            NOTIONAL_AMT = 0 # 合约名义本金
            try:
                I_CODE = trade_order['I_CODE_BLOOMBERG']
                I_NAME = trade_order['I_NAME']
                LAST_TRADE_DATE = trade_order['LAST_TRADE_DATE']
                if TRADE_FINISH_DATE:
                    if LAST_TRADE_DATE > TRADE_FINISH_DATE:
                        TRADE_FINISH_DATE = LAST_TRADE_DATE
                else:
                    TRADE_FINISH_DATE = LAST_TRADE_DATE
            except:
                pass
            OPEN_DEACOUNT = trade_order['OPEN_DEACOUNT'] # 开仓名义数量
            TRGT_CHOICES = {'10':'多','11':'空','20':'空','21':'多'}
            TRADE_TYPE = TRGT_CHOICES.get(trade_order['TRADE_TYPE']) 
            if TRADE_TYPE == '多':
                NOTIONAL_AMT = trade_order['DEAAMOUNT'] + trade_order['DEAFEE']
            elif TRADE_TYPE == '空':
                NOTIONAL_AMT = trade_order['DEAAMOUNT'] - trade_order['DEAFEE']
            NOTIONAL_AMT_TOTAL += NOTIONAL_AMT
            
            MULTIPLIER = trade_order['MULTIPLIER'] # 合约乘数
            AVGCOST = trade_order['AVGCOST'] # 期初价格
            trade_info = {
                'NAME': I_NAME, # 标的合约	
                'CODE': I_CODE, # 标的代码
                'M_TYPE': trade_order['M_TYPE'],
                'CAP_RATE': str(round(float(trade_order['CAP_RATE']) * 100, 6)) + "%",
                'CAP_LOC': trade_order['CAP_LOC'],
                'IM_RATE': str(round(float(trade_order['IM_RATE']) * 100, 6)) + "%",
                'VM_RATE': str(round(float(trade_order['VM_RATE']) * 100, 6)) + "%",
                'BASIC_FEE_RATE': str(round(float(trade_order['BASIC_FEE_RATE']) * 100, 6)) + "%",
                'AVGCOST': str(round(AVGCOST,6)), # 标的期初价格（交易货币）
                'DEACOUNT': str(OPEN_DEACOUNT), # 标的名义数量	
                'MULTIPLIER': str(MULTIPLIER), # 合约乘数	
                'TRADE_TYPE':TRADE_TYPE, # 方向
                'LAST_TRADE_DATE': LAST_TRADE_DATE, # 最后平仓日
            }
            trade_ordr_lst.append(trade_info)
        trade_summary['TRADE_DEAL_DATE'] = TRADE_DEAL_DATE
        trade_summary['TRADE_ORDER_DATE'] = TRADE_ORDER_DATE
        trade_summary['TRADE_FINISH_DATE'] = TRADE_FINISH_DATE
        trade_summary['NOTIONAL_AMT_TOTAL'] = str(NOTIONAL_AMT_TOTAL)
        trade_summary['SIGN_DATE'] = SIGN_DATE
        return {'trade_ordr_lst':trade_ordr_lst, 'trade_summary':trade_summary}

    def trade_unordr_handle(self, data):
        trade_ordr_lst = []
        trade_summary = {}

        SIGN_DATE = "" # 日期
        if not data.empty:
            for _idx, row in data.iterrows():
                if _idx == 0:
                    TRADE_ORDER_DATE = row['CLOSE_DATE']
                    trade_date_lst = TRADE_ORDER_DATE.split('-')
                    if trade_date_lst:
                        SIGN_DATE = trade_date_lst[0] + ' 年 ' + trade_date_lst[1] + ' 月 ' + trade_date_lst[2] + ' 日'
                ORDER_CFM_DOC_ID_id = row['OPEN_ORDER_CFM_DOC_ID']
                TRGT_CHOICES = {'10':'多','11':'空','20':'空','21':'多'}
                TRADE_TYPE = TRGT_CHOICES.get(row['TRADE_TYPE']) 
                TRADE_DEAL_DATE = row['CLOSE_DATE'].replace('-', '/')
                NAME = row['I_NAME']
                CODE = row['I_CODE_BLOOMBERG']
                CLOSE_DEAAMOUNT = row['CLOSE_DEAAMOUNT']
                CLOSE_DEAFEE = row['CLOSE_DEAFEE']
                DEACOST_POS = row['DEACOST_POS']
                DEACOUNT_POS = row['DEACOUNT_POS']
                DEACOUNT = row['DEACOUNT']
                DEACOUNT_TOTAL = 0
                LAST_DEAAMOUNT = 0
                if TRADE_TYPE == '多':
                    DEACOUNT_TOTAL = DEACOUNT_POS + DEACOUNT
                    LAST_DEAAMOUNT = CLOSE_DEAAMOUNT - CLOSE_DEAFEE
                elif TRADE_TYPE == '空':
                    DEACOUNT_TOTAL = DEACOUNT_POS - DEACOUNT
                    LAST_DEAAMOUNT = CLOSE_DEAAMOUNT + CLOSE_DEAFEE
                trade_info = {
                    'CFM_DOC_ID': ORDER_CFM_DOC_ID_id, # 交易确认书编号
                    'TRADE_TYPE': TRADE_TYPE, # 交易方向
                    'TRADE_DEAL_DATE': TRADE_DEAL_DATE, # 平仓确认日期
                    'NAME': NAME, # 标的名称
                    'CODE': CODE, # 标的代码
                    'DACTB':str(DEACOUNT_TOTAL), # 平仓前标的合约名义数量
                    'DACT':str(DEACOUNT), # 平仓成交数量
                    'DACTA':str(DEACOUNT_POS), # 平仓后标的合约名义数量
                    'AMT': str(LAST_DEAAMOUNT), # 平仓成交金额（交易货币）
                    'DEACOST': str(DEACOST_POS) # 平仓后标的合约名义本金额（交易货币）
                }
                trade_ordr_lst.append(trade_info)
        trade_summary['SIGN_DATE'] = SIGN_DATE
        return {'trade_ordr_lst':trade_ordr_lst, 'trade_summary':trade_summary}

    def create_qrcode_handel(self):
        # 为参数对象创建二维码图片
        pass
ConfirmAttGenCtrlObj = ConfirmAttGenCtrl()
